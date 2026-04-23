# This script scans a folder, extracts text from various file types (PDF, DOCX, TXT, etc.),
# performs OCR on image-based PDFs, cleans the text, and anonymizes PII using
# Microsoft Presidio (NLP-backed detection for names, emails, phones, fiscal codes,
# license plates, IBANs, credit cards, etc.).
# It uses parallel processing for greater efficiency.

import os
import sys
import json
import re
import uuid
import hashlib
import logging
from pathlib import Path
from concurrent.futures import ProcessPoolExecutor, as_completed
from tqdm import tqdm
from typing import Tuple, Optional
from datetime import datetime

from database import get_session, create_db_and_tables
from models import Document
from sqlmodel import select
from config_manager import get_config

import pytesseract
import fitz
import docx
from PIL import Image
import xlrd
from striprtf.striprtf import rtf_to_text
from odf.opendocument import load as odf_load
from odf import text as odf_text, teletype as odf_teletype
import xml.etree.ElementTree as ET

# --- GLOBAL CONFIGURATION AND PATHS ---
SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent

extract_config = get_config('extract_text.py', {})
raw_folder = get_config('input_folder', '')
FOLDER_TO_SCAN = (PROJECT_ROOT / raw_folder) if raw_folder.strip() else None

OUTPUT_FOLDER = PROJECT_ROOT / get_config('database_folder', './data/output')
LOG_FOLDER = PROJECT_ROOT / get_config('log_folder', './logs')

TESSERACT_PATH = extract_config.get('tesseract_path', '')
MAX_WORKERS = extract_config.get('max_workers', os.cpu_count() or 4)

if TESSERACT_PATH:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
else:
    logging.warning("Tesseract path (TESSERACT_PATH) not specified; image PDF OCR might not work.")

from utils_logging import setup_logging
from utils_presidio import create_analyzer, create_anonymizer, anonymize_text

# --- PER-WORKER PRESIDIO ENGINE GLOBALS ---
# These are initialized inside each worker process via worker_init().
# spaCy models cannot be pickled across process boundaries, so each
# worker must create its own AnalyzerEngine and AnonymizerEngine.
_worker_analyzer = None
_worker_anonymizer = None

# --- GENERAL UTILITY FUNCTIONS ---
def clean_text(text: str) -> str:
    """Cleans text from unwanted characters and normalizes spaces."""
    if not text: return ""
    text = re.sub(r'[^\x20-\x7E\n\r\t]', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def generate_md5_hash(file_path: Path) -> Optional[str]:
    """Calculates the MD5 hash (digital fingerprint) of a file."""
    h = hashlib.md5()
    try:
        with file_path.open('rb') as f:
            for chunk in iter(lambda: f.read(65536), b''):
                h.update(chunk)
        return h.hexdigest()
    except IOError as e:
        logging.error(f"I/O Error calculating hash for {file_path.name}: {e}")
        return None
    except Exception as e:
        logging.error(f"Unexpected error calculating hash for {file_path.name}: {e}")
        return None

def get_scaling_factor(page_rect_width: float, page_rect_height: float,
                       default_dpi=300, max_img_width=10000, max_img_height=10000) -> float:
    """Calculates scaling factor for PDF images to use in OCR."""
    factor = default_dpi / 72.0
    current_width = page_rect_width * factor
    current_height = page_rect_height * factor
    if current_width > max_img_width or current_height > max_img_height:
        scale_x = max_img_width / page_rect_width
        scale_y = max_img_height / page_rect_height
        factor = min(scale_x, scale_y)
    return factor

# --- TEXT EXTRACTION ENGINES BY FILE TYPE ---
def extract_text_from_pdf(file_path: Path) -> str:
    """Extracts text from PDF files, with OCR fallback for image pages."""
    extracted_pages = []
    ocr_triggered_on_doc = False
    max_ocr_pages_limit = 2
    doc = None

    try:
        if not file_path.exists():
            logging.error(f"PDF file not found (Path.exists check failed): {file_path.resolve()}")
            return ""
        if not os.access(file_path, os.R_OK):
            logging.error(f"Read permission missing for PDF file (os.access check failed): {file_path.resolve()}")
            return ""

        doc = fitz.open(str(file_path)) 
        for i, page in enumerate(doc):
            if i >= 20 and not ocr_triggered_on_doc:
                 logging.info(f"PDF {file_path.name}: Reading interrupted after 20 textual pages.")
                 break
            page_text = page.get_text("text").strip()
            if not page_text and i < max_ocr_pages_limit:
                ocr_triggered_on_doc = True
                logging.info(f"PDF {file_path.name}, page {i+1}: empty text, starting OCR.")
                try:
                    factor = get_scaling_factor(page.rect.width, page.rect.height)
                    pix = page.get_pixmap(matrix=fitz.Matrix(factor, factor))
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    page_text = pytesseract.image_to_string(img, lang='ita+eng').strip()
                    logging.info(f"PDF {file_path.name}, page {i+1}: OCR completed.")
                except Exception as ocr_e:
                    logging.error(f"OCR Error PDF {file_path.name}, page {i+1}: {ocr_e}")
                    page_text = ""
            extracted_pages.append(page_text)
            if ocr_triggered_on_doc and (i + 1) >= max_ocr_pages_limit and (i + 1) < len(doc):
                logging.info(f"PDF {file_path.name}: Reading interrupted after {max_ocr_pages_limit} OCR pages.")
                break
        full_text = "\n---\n".join(p for p in extracted_pages if p)
        return clean_text(full_text)

    except RuntimeError as fitz_runtime_error: 
        error_message = str(fitz_runtime_error)
        if "Failed to open file" in error_message:
            logging.error(f"PyMuPDF (fitz) failed to open the PDF file: '{file_path.resolve()}'. Message: '{error_message}'.")
        else:
            logging.error(f"PyMuPDF (fitz) runtime error processing '{file_path.name}': {error_message}")
        return ""
    except Exception as e: 
        logging.error(f"Unexpected generic error extracting from PDF '{file_path.name}': {e} (Type: {type(e).__name__})")
        return ""
    finally:
        if doc: 
            doc.close()

def extract_text_from_docx(file_path: Path) -> str:
    """Extracts text from DOCX (Word) files."""
    try:
        doc = docx.Document(str(file_path))
        return clean_text("\n".join(p.text for p in doc.paragraphs if p.text))
    except Exception as e:
        logging.error(f"Error extracting DOCX {file_path.name}: {e}")
        return ""

def extract_text_from_doc(file_path: Path) -> str:
    """Extracts text from DOC files (legacy Word, Windows only)."""
    if sys.platform != "win32":
        logging.warning(f"Extracting .doc is supported on Windows only. Skipped: {file_path.name}")
        return ""
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()  # FIX 2: Initialize COM for multi-processing threads
    except ImportError:
        logging.error("pywin32 library not installed; impossible to process .doc.")
        return ""
    
    word_app = None
    doc_com = None
    try:
        word_app = win32com.client.Dispatch("Word.Application")
        word_app.Visible = False
        doc_com = word_app.Documents.Open(str(file_path), ReadOnly=True)
        text = doc_com.Content.Text
        return clean_text(text)
    except Exception as e:
        logging.error(f"Error extracting DOC {file_path.name} (win32com): {e}")
        return ""
    finally:
        if doc_com: doc_com.Close(False)
        if word_app: word_app.Quit()
        try:
            import pythoncom
            pythoncom.CoUninitialize() # Clean up COM
        except Exception:
            pass

def extract_text_from_odt(file_path: Path) -> str:
    """Extracts text from ODT (OpenDocument Text) files."""
    try:
        doc = odf_load(str(file_path))
        paragraphs = doc.getElementsByType(odf_text.P)
        return clean_text("\n".join(odf_teletype.extractText(p) for p in paragraphs))
    except Exception as e:
        logging.error(f"Error extracting ODT {file_path.name}: {e}")
        return ""

def extract_text_from_rtf(file_path: Path) -> str:
    """Extracts text from RTF files."""
    try:
        raw_text = file_path.read_text(encoding='utf-8', errors='ignore')
        return clean_text(rtf_to_text(raw_text))
    except Exception as e:
        logging.error(f"Error extracting RTF {file_path.name}: {e}")
        return ""

def extract_text_from_xml(file_path: Path) -> str:
    """Extracts text from XML files (uses structured parsing, regex fallback)."""
    try:
        tree = ET.parse(str(file_path))
        root = tree.getroot()
        text_content = "".join(root.itertext())
        return clean_text(text_content)
    except ET.ParseError as e:
        logging.warning(f"XML parsing error {file_path.name}: {e}. Attempting with regex.")
        try:
            raw_text = file_path.read_text(encoding='utf-8', errors='ignore')
            return clean_text(re.sub(r'<[^>]+>', '', raw_text))
        except Exception as fallback_e:
            logging.error(f"XML extraction error (regex fallback) {file_path.name}: {fallback_e}")
            return ""
    except Exception as e:
        logging.error(f"Unexpected XML extraction error {file_path.name}: {e}")
        return ""

def _extract_text_from_json_recursive(data) -> list:
    """Recursive helper function to extract all strings from JSON data."""
    text_parts = []
    if isinstance(data, dict):
        for value in data.values():
            text_parts.extend(_extract_text_from_json_recursive(value))
    elif isinstance(data, list):
        for item in data:
            text_parts.extend(_extract_text_from_json_recursive(item))
    elif isinstance(data, str):
        text_parts.append(data)
    return text_parts

def extract_text_from_json(file_path: Path) -> str:
    """Extracts all text values (strings) from a JSON file."""
    try:
        raw_json = file_path.read_text(encoding='utf-8', errors='ignore')
        data = json.loads(raw_json)
        all_strings = _extract_text_from_json_recursive(data)
        return clean_text(" ".join(all_strings))
    except json.JSONDecodeError as e:
        logging.error(f"JSON decoding error {file_path.name}: {e}")
        return ""
    except Exception as e:
        logging.error(f"JSON extraction error {file_path.name}: {e}")
        return ""

def extract_text_from_html(file_path: Path) -> str:
    """Extracts text from HTML files using regex."""
    try:
        raw_html = file_path.read_text(encoding='utf-8', errors='ignore')
        raw_html = re.sub(r'<(script|style)\b[^>]*>.*?</\1\s*>', '', raw_html, flags=re.IGNORECASE | re.DOTALL)
        text_no_tags = re.sub(r'<[^>]+>', '', raw_html)
        return clean_text(text_no_tags)
    except Exception as e:
        logging.error(f"HTML extraction error {file_path.name}: {e}")
        return ""

def extract_text_from_csv(file_path: Path) -> str:
    """Extracts text from CSV files (treated as plain text)."""
    try:
        lines = file_path.read_text(encoding='utf-8', errors='ignore').splitlines()
        return clean_text("\n".join(lines))
    except Exception as e:
        logging.error(f"CSV extraction error {file_path.name}: {e}")
        return ""

def extract_text_from_xls(file_path: Path) -> str:
    """Extracts text from XLS files (legacy Excel)."""
    try:
        with open(os.devnull, 'w') as devnull:
            workbook = xlrd.open_workbook(str(file_path), logfile=devnull)
        lines = []
        for sheet in workbook.sheets():
            for r_idx in range(sheet.nrows):
                row_values = [str(cell.value) if cell.ctype != xlrd.XL_CELL_EMPTY else ""
                              for cell in sheet.row(r_idx)]
                lines.append(" ".join(row_values).strip())
        return clean_text("\n".join(filter(None, lines)))
    except Exception as e:
        logging.error(f"XLS extraction error {file_path.name}: {e}")
        return ""

def extract_text_from_plaintext(file_path: Path) -> str:
    """Extracts text from plain text files (TXT, MD)."""
    try:
        return clean_text(file_path.read_text(encoding='utf-8', errors='ignore'))
    except Exception as e:
        logging.error(f"Plaintext extraction error from {file_path.name}: {e}")
        return ""

# --- Extension -> Extraction Function Map ---
EXTRACTOR_MAP = {
    '.pdf': extract_text_from_pdf, '.docx': extract_text_from_docx,
    '.doc': extract_text_from_doc, '.odt': extract_text_from_odt,
    '.rtf': extract_text_from_rtf, '.xml': extract_text_from_xml,
    '.json': extract_text_from_json, '.html': extract_text_from_html,
    '.htm': extract_text_from_html, '.csv': extract_text_from_csv,
    '.xls': extract_text_from_xls, '.md': extract_text_from_plaintext,
    '.txt': extract_text_from_plaintext,
}

# --- SINGLE FILE PROCESSING PIPELINE ---
def process_file(file_path: Path) -> dict | None:
    """Processes a single file: calculates hash, extracts text, anonymizes PII via Presidio."""
    global _worker_analyzer, _worker_anonymizer
    logging.debug(f"Starting processing: {file_path.name}")
    file_hash = generate_md5_hash(file_path)
    if not file_hash:
        logging.warning(f"Hash cannot be calculated for {file_path.name}, file skipped.")
        return None

    ext = file_path.suffix.lower()
    extractor_func = EXTRACTOR_MAP.get(ext)
    extracted_text = ""
    if extractor_func:
        extracted_text = extractor_func(file_path)
    else:
        logging.info(f"No extractor for format '{ext}' of {file_path.name}. Skipped.")
        return None

    masked_text, pii_flag = anonymize_text(extracted_text, _worker_analyzer, _worker_anonymizer)
    
    file_type = "File"
    parent_id = None
    if file_path.name.startswith("email_"):
        if "_att_" in file_path.name:
            file_type = "Attachment"
            # Pattern: email_{msg_id}_att_{counter}_{name}
            match = re.search(r"^email_([^_]+)_att_", file_path.name)
            if match:
                parent_id = match.group(1)
        elif file_path.name.endswith("_body.txt"):
            file_type = "Email"
            # Pattern: email_{msg_id}_body.txt
            match = re.search(r"^email_([^_]+)_body\.txt$", file_path.name)
            if match:
                parent_id = match.group(1)

    # Fetch the OS-level modification timestamp (which we spoof in Step 0 for emails)
    try:
        creation_timestamp = os.path.getmtime(file_path)
        creation_date = datetime.fromtimestamp(creation_timestamp)
    except Exception as e:
        logging.warning(f"Could not retrieve creation date for {file_path.name}: {e}")
        creation_date = None

    record = {
        'id': str(uuid.uuid4()),
        'type': file_type,
        'parent_id': parent_id,
        'file_path': str(file_path.resolve()),
        'file_name': file_path.name,
        'extracted_text_masked': masked_text,
        'md5_hash': file_hash,
        'names_or_surnames_masked': pii_flag,
        'creation_date': creation_date
    }
    logging.debug(f"FileID: {record['id']} - Processed: {file_path.name}")
    return record

def worker_init(log_dir: Path):
    """Initializes logging and Presidio NLP engine in each worker process."""
    # Each worker needs its own engine because spaCy models can't be pickled.
    global _worker_analyzer, _worker_anonymizer
    setup_logging("1_extract_text")
    _worker_analyzer = create_analyzer()
    _worker_anonymizer = create_anonymizer()
    logging.info("Worker Presidio engine initialized.")

# --- DIRECTORY SCANNING AND PARALLEL EXECUTION ---
def process_files_in_directory(
    scan_directory: Path, output_dir: Path, log_dir: Path
) -> None:
    """Scans a directory, processes supported files in parallel, saves results to JSON."""
    if not scan_directory.is_dir():
        logging.error(f"Folder to scan '{scan_directory}' does not exist.")
        sys.exit(1)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    output_file = output_dir / "extracted_documents.json"

    all_files = [f for f in scan_directory.glob("**/*") if f.is_file()]
    supported_extensions = set(EXTRACTOR_MAP.keys())
    files_to_process = [f for f in all_files if f.suffix.lower() in supported_extensions]
    
    skipped_files_count = len(all_files) - len(files_to_process)
    if skipped_files_count > 0:
        logging.info(f"Found {len(all_files)} files; {skipped_files_count} with unsupported extensions will be skipped.")
    logging.info(f"Starting processing of {len(files_to_process)} supported files from {scan_directory}")

    create_db_and_tables()
    results = []
    hashes_seen = set()

    try:
        with get_session() as session:
            existing_docs = session.exec(select(Document.md5_hash)).all()
            hashes_seen.update(existing_docs)
        logging.info(f"Loaded {len(hashes_seen)} existing hashes from database. New files will be appended safely.")
    except Exception as e:
        logging.warning(f"Could not read existing database, starting fresh. Error: {e}")
    
    # ProcessPoolExecutor with initializer to configure logging in workers
    with ProcessPoolExecutor(max_workers=MAX_WORKERS, initializer=worker_init, initargs=(log_dir,)) as executor:
        futures = [executor.submit(process_file, f) for f in files_to_process]
        for future in tqdm(as_completed(futures), total=len(futures), desc="Processing files"):
            try:
                res = future.result()
                if res:
                    if res['md5_hash'] not in hashes_seen:
                        hashes_seen.add(res['md5_hash'])
                        results.append(res)
                    else:
                        logging.info(f"Duplicate file (MD5) skipped: {res['file_path']}")
            except Exception as exc:
                logging.error(f"Error during parallel processing of a file: {exc}")
                
    try:
        if results:
            with get_session() as session:
                docs_to_insert = [Document(**res) for res in results]
                session.add_all(docs_to_insert)
                session.commit()
        logging.info(f"Saving completed. {len(results)} new unique records inserted into Database.")

        # Delete anonymized files if we downloaded them from Gmail
        active_source = get_config('active_source', 'local')
        mail_config = get_config('0_extract_mail.py', {})
        if active_source in ('gmail', 'outlook') and mail_config.get('delete_after_processing', True):
            for res in results:
                file_path = Path(res['file_path'])
                try:
                    if file_path.exists():
                        file_path.unlink()
                except Exception as del_err:
                    logging.error(f"Failed to delete anonymized file {file_path}: {del_err}")
            logging.info("Deleted processed Gmail original files successfully.")

    except IOError as e:
        logging.error(f"Error writing output JSON file '{output_file}': {e}")
    except Exception as e:
        logging.error(f"Unexpected error saving JSON results: {e}")

# --- MAIN SCRIPT EXECUTION BLOCK ---
if __name__ == "__main__":
    setup_logging("1_extract_text") 
    logging.info("Logging system configured for Main process.")

    # FIX 3: Correct empty folder validation
    if not FOLDER_TO_SCAN:
        logging.error("'input_folder' path not specified or empty in the configuration database.")
        sys.exit(1)

    logging.info(f"Starting script. Folder to scan: {FOLDER_TO_SCAN}")
    logging.info(f"Output folder: {OUTPUT_FOLDER}")
    logging.info(f"Logs folder: {LOG_FOLDER}")
    logging.info(f"Maximum number of parallel workers: {MAX_WORKERS}")

    process_files_in_directory(
        FOLDER_TO_SCAN,
        OUTPUT_FOLDER,
        LOG_FOLDER
    )
    logging.info("Script completed.")